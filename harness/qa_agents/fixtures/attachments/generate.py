"""Regenerate the attachment fixture pack used by Cătălina's persona
and by any harness test that needs real, valid files in their format.

Run from anywhere in the repo:

    uv run --with reportlab --with python-docx --with openpyxl --with Pillow \\
        qa-agents/harness/qa_agents/fixtures/attachments/generate.py

These fixtures are deliberately small, hand-crafted and synthetic —
they ARE real files in their format (a valid PDF is a valid PDF, a
docx is a real Open XML zip, an xlsx is a real Open XML zip) so the
AI providers ingest them correctly. Content is fictional so there's
no embedded PII and we know exactly what's in each file. See README.md
in this directory for what each fixture is for.
"""
from __future__ import annotations

import os
import zipfile
from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

HERE = Path(__file__).parent


def _invoice_pdf() -> None:
    """Small valid PDF (~5–10 KB) with realistic-looking invoice text."""
    out = HERE / "sample-invoice.pdf"
    c = canvas.Canvas(str(out), pagesize=A4)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, 800, "Doyle & Sons Plumbing Ltd")
    c.setFont("Helvetica", 10)
    c.drawString(72, 782, "12 Holmwood Lane, Sheffield S10 4AB")
    c.drawString(72, 768, "VAT no. GB 123 4567 89")
    c.setFont("Helvetica-Bold", 11)
    c.drawString(72, 740, "Invoice #INV-2026-014")
    c.setFont("Helvetica", 10)
    c.drawString(72, 724, "Date: 22 May 2026")
    c.drawString(72, 708, "Bill to: Cat Popa, Lambert & Co Accountancy")
    c.setFont("Helvetica-Bold", 10)
    c.drawString(72, 670, "Description")
    c.drawString(330, 670, "Net")
    c.drawString(400, 670, "VAT")
    c.drawString(470, 670, "Total")
    c.setFont("Helvetica", 10)
    c.drawString(72, 652, "Boiler annual service")
    c.drawString(330, 652, "£120.00")
    c.drawString(400, 652, "£24.00")
    c.drawString(470, 652, "£144.00")
    c.drawString(72, 636, "Pressure-relief valve replacement")
    c.drawString(330, 636, "£45.00")
    c.drawString(400, 636, "£9.00")
    c.drawString(470, 636, "£54.00")
    c.line(72, 620, 540, 620)
    c.drawString(330, 604, "Subtotal")
    c.drawString(470, 604, "£165.00")
    c.drawString(330, 588, "VAT @ 20%")
    c.drawString(470, 588, "£33.00")
    c.setFont("Helvetica-Bold", 11)
    c.drawString(330, 568, "Total due")
    c.drawString(470, 568, "£198.00")
    c.setFont("Helvetica", 9)
    c.drawString(72, 480, "Payment terms: net 30 days. Thank you for your business.")
    c.showPage()
    c.save()


def _invoice_large_pdf() -> None:
    """~6 MB valid PDF for tier-cap testing.

    Embeds a 1000×1000 random-RGB PNG eight times — random bytes don't
    compress, so the embedded image is ~3 MB and the resulting PDF is
    in the 6–8 MB range. Still a real, provider-ingestible PDF — not a
    padded placeholder.
    """
    img_path = HERE / "_pad.png"
    Image.frombytes(
        "RGB", (1000, 1000), os.urandom(3 * 1000 * 1000)
    ).save(img_path, "PNG", optimize=False)
    try:
        out = HERE / "sample-invoice-large.pdf"
        c = canvas.Canvas(str(out), pagesize=A4)
        for i in range(2):
            c.setFont("Helvetica-Bold", 14)
            c.drawString(72, 800, f"Doyle & Sons — bulk-attachment fixture page {i + 1}")
            c.setFont("Helvetica", 9)
            c.drawString(
                72,
                782,
                "Synthetic fixture (qa-agents persona harness) — random-noise image"
                " padding for tier-cap testing.",
            )
            c.drawImage(str(img_path), 72, 200, width=450, height=450)
            c.showPage()
        c.save()
    finally:
        img_path.unlink(missing_ok=True)


def _report_docx() -> None:
    """Valid DOCX (~30 KB) with realistic accountant content."""
    out = HERE / "sample-report.docx"
    d = Document()
    d.add_heading("Quarterly Bookkeeping Summary — Q1 2026", 0)
    d.add_paragraph(
        "Prepared by Cătălina Popa for Lambert & Co Accountancy. "
        "This document is a synthetic fixture used by SlyReply's QA "
        "persona harness — see qa-agents/PERSONAS.md §12."
    )
    d.add_heading("Receivables", 1)
    d.add_paragraph("Total invoiced in Q1: £124,830.00.")
    d.add_paragraph("Days sales outstanding: 38 days, up from 31 in Q4 2025.")
    d.add_heading("Payables", 1)
    d.add_paragraph("Total paid in Q1: £58,420.00.")
    d.add_paragraph(
        "Largest single payable: HMRC PAYE for January (£18,420.00). "
        "Settled on the due date."
    )
    d.add_heading("Net position", 1)
    d.add_paragraph("Net Q1 cash flow: +£66,410.00.")
    d.save(str(out))


def _figures_xlsx() -> None:
    """Valid XLSX (~5 KB) with a small Q1 table."""
    out = HERE / "sample-figures.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Q1 2026"
    ws.append(["Month", "Invoiced (£)", "Paid (£)", "Net (£)"])
    ws.append(["January", 42500, 19800, 22700])
    ws.append(["February", 39800, 22100, 17700])
    ws.append(["March", 42530, 16520, 26010])
    ws.append(["Total", 124830, 58420, 66410])
    wb.save(str(out))


def _expenses_csv() -> None:
    """Trivial CSV — plain text, half a dozen rows."""
    out = HERE / "sample-expenses.csv"
    rows = [
        "date,description,category,amount_gbp,vat_gbp",
        "2026-01-04,Software subscription,Software,49.00,9.80",
        "2026-01-12,Train to client meeting,Travel,82.50,0.00",
        "2026-02-03,Office supplies,Stationery,38.40,7.68",
        "2026-02-18,Accountancy CPD course,Training,295.00,59.00",
        "2026-03-09,Domain renewal,Software,12.00,2.40",
    ]
    out.write_text("\n".join(rows) + "\n", encoding="utf-8")


def _receipt_png() -> None:
    """Small synthetic 'receipt' image — white background, black text."""
    out = HERE / "sample-receipt.png"
    img = Image.new("RGB", (300, 360), "white")
    draw = ImageDraw.Draw(img)
    lines = [
        "STOCKPORT CONVENIENCE",
        "12 High Street, Stockport",
        "",
        "Receipt #4729  22-May-2026",
        "",
        "Tea (3 boxes)        7.50",
        "Biscuits             1.80",
        "VAT @ 20%            1.86",
        "",
        "Total              11.16",
        "",
        "Thank you. Please come again.",
    ]
    y = 20
    for line in lines:
        draw.text((20, y), line, fill="black")
        y += 22
    img.save(out, "PNG")


def _receipt_jpg() -> None:
    """Same content as the PNG, JPEG-encoded."""
    png = HERE / "sample-receipt.png"
    if not png.exists():
        _receipt_png()
    out = HERE / "sample-receipt.jpg"
    Image.open(png).convert("RGB").save(out, "JPEG", quality=85)


def _zero_byte() -> None:
    """Edge case: a legitimately empty file."""
    (HERE / "sample-empty.txt").write_bytes(b"")


def _zip_of_pdfs() -> None:
    """Edge case: a ZIP containing a couple of small PDFs.

    Tests how the pipeline handles compound types — does it treat the
    zip as one unsupported file, unpack it, or something else?
    """
    pdf = HERE / "sample-invoice.pdf"
    if not pdf.exists():
        _invoice_pdf()
    out = HERE / "sample-bundle.zip"
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.write(pdf, arcname="invoice-014.pdf")
        zf.write(pdf, arcname="invoice-015.pdf")
    out.write_bytes(buf.getvalue())


def _renamed_docx_as_pdf() -> None:
    """Edge case: a DOCX renamed to .pdf — content-type vs extension mismatch.

    Tests whether the pipeline trusts the filename or sniffs the actual
    bytes. Important because real-world users do this all the time.
    """
    docx = HERE / "sample-report.docx"
    if not docx.exists():
        _report_docx()
    (HERE / "sample-mislabeled.pdf").write_bytes(docx.read_bytes())


def main() -> None:
    _invoice_pdf()
    _invoice_large_pdf()
    _report_docx()
    _figures_xlsx()
    _expenses_csv()
    _receipt_png()
    _receipt_jpg()
    _zero_byte()
    _zip_of_pdfs()
    _renamed_docx_as_pdf()
    print("Fixture pack regenerated in", HERE)


if __name__ == "__main__":
    main()
